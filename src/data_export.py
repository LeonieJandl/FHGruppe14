from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from matplotlib.figure import Figure


def create_report_template():
    # Create new document
    doc = Document()

    # Add title
    title = doc.add_heading(level=1)
    title_run = title.add_run("Bericht [week]")
    title_run.font.size = Pt(16)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add sections with placeholder text
    sections = [
        "Wochenbericht",
        "Vergleich zur vorherigen Woche",
    ]

    for section in sections:
        doc.add_heading(section, level=2)
        doc.add_paragraph(f"[{section.lower().replace(" ", "")}]")
        # Add placeholder for potential images in each section
    doc.add_paragraph("[image]").alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save the document
    doc.save("./templates/report_template.docx")


def insert_content(
    content_dict: dict[str, str],
    image: Figure,
    template_path: Optional[str] = "./templates/report_template.docx",
):
    """
    Insert actual content into the template

    content_dict should be a dictionary with keys matching section names
    and values containing the text and image paths
    """
    doc = Document(template_path)

    # Replace the week placeholder in title
    for paragraph in doc.paragraphs:
        if "[week]" in paragraph.text:
            for run in paragraph.runs:
                run.text = run.text.replace("[week]", content_dict.get("week", ""))

    # Replace content placeholders
    for paragraph in doc.paragraphs:
        if "[wochenbericht]" in paragraph.text:
            paragraph.text = content_dict.get("Wochenbericht", "")
        elif "[vergleichzurvorherigenwoche]" in paragraph.text:
            paragraph.text = content_dict.get("Vergleich zur vorherigen Woche", "")

        # Replace image placeholder with actual image
        if "[image]" in paragraph.text:
            paragraph.text = ""  # Clear the placeholder text
            # Save the matplotlib figure to a temporary file
            image.savefig("temp_plot.png")
            # Add the image to the document
            doc.add_picture("temp_plot.png", width=Inches(6))  # Adjust width as needed
            # Remove temporary file
            import os

            os.remove("temp_plot.png")

    doc.save("report_final.docx")


if __name__ == "__main__":
    # Create the template
    create_report_template()
